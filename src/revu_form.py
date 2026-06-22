"""
revu_form.py — 레뷰(REVU) 네이버 베이직 체험단 양식 docx 생성.

빈 표준 양식(templates/revu_basic_template.docx)을 원본으로 로드하여, 표의 정해진
빈칸 셀에만 입력값을 채워 넣은 docx 를 만든다. 새로 그리지 않는다 — 원본 복사 후
대상 셀의 텍스트만 치환하므로 서식·사전안내문·표 구조가 100% 보존된다.

★빈 양식 표 구조(python-docx 분석 결과, 표 인덱스 0~12):
  표0  r1.p0  "2. 콘텐츠 타입 선택:  "      → 콘텐츠 타입(블로그/클립)
  표3  r0.p0  "성함: "                        → 담당자 성함
       r0.p1  "연락처: "                      → 담당자 연락처
       r0.p2  "이메일: "                       → 담당자 이메일
  표4  r0.p0  "[      ] "                      → 캠페인 제목(20자)
  표5  r0.p0  "[      ] "                      → 캠페인 부제목(40자)
  표6  r0.p0  "　　명"                         → 모집인원
  표7  r0.p0  "제품명: "                       → 제품명
       r0.p1  "제공수량:   개"                 → 제공수량(원본 "개" 제거, 입력값 그대로)
  표8  r1.p0~2 "1. " "2. " "3. "  (블로그 미션) → 미션(콘텐츠=블로그일 때)
       r3.p0~2 "1. " "2. " "3. "  (클립 미션)   → 미션(콘텐츠=클립일 때)
  표10 r1.p0  "제목키워드 (1~3개) : "          → 제목키워드
       r1.p1  "본문키워드 (3~5개) : "          → 본문키워드
  표11 r0.p0  "링크입력: "                     → 제품링크 URL
  표12 (29행 사전 안내문) — ★고정. 절대 건드리지 않는다.

표1·2(이미지 안내), 표9(구매평 미션), 표8 상품ID/해시태그 블록은 양식 기본값 그대로 둔다.
"""

from __future__ import annotations

import io
import json
import pathlib
import re
from dataclasses import dataclass, field

# ★python-docx(docx)는 docx 생성 시점에만 필요하므로 ★지연 import 한다(build_revu_docx 내부).
#   모듈 최상단에서 import 하면 배포본에 python-docx 가 없을 때 이 모듈 전체가 import 불가가 되어,
#   find_banned_words·merge_keywords·추적URL 검증 등 순수 함수와 앱 전체(app.py line 76)가
#   함께 죽는다. 지연 import 로 docx 미설치 시에도 양식 외 기능은 정상 동작한다.

# 템플릿 경로(프로젝트 루트 기준).
TEMPLATE_PATH = pathlib.Path(__file__).resolve().parent.parent / "templates" / "revu_basic_template.docx"

# 글자수 제한(양식 규칙).
TITLE_MAX = 20
SUBTITLE_MAX = 40

# 자동 채움 기본값(담당자 — 화면에서 수정 가능).
DEFAULT_MANAGER_NAME = "박민우"
DEFAULT_MANAGER_PHONE = "010-3924-1155"
DEFAULT_MANAGER_EMAIL = "dgkorea93@naver.com"

# ── 네이버 유입 추적 URL(nt_) 조립 ───────────────────────────────────────────
# 네이버 공식 규칙:
#  · nt_source/nt_medium : 영문·숫자·특수문자 3종(-,_,.)만. 한글·공백 금지. 필수.
#  · nt_detail/nt_keyword: 한글·영문·숫자·특수문자 3종(-,_,.) 허용. 공백 금지. 선택.
#  · 그 외 특수문자(/,?,&,=,# 등)·공백은 모두 금지.
DEFAULT_NT_SOURCE = "naver.blog"
DEFAULT_NT_MEDIUM = "social"
DEFAULT_NT_DETAIL = "revu"

# ── 양식 저장/불러오기(JSON 파일) ────────────────────────────────────────────
# ★Streamlit Cloud 영구저장 안 씀 — 시경 PC 로의 파일 다운로드/업로드만(앱 재시작 무관).
# 화면 위젯 key → 기본값. 저장은 이 키들의 session_state 값을 JSON 으로, 불러오기는
# JSON 을 이 키들의 session_state 로 되돌린다(위젯 key 기반 정확 복원).
FORM_VERSION = 1

REVU_SAVE_FIELDS: list[tuple[str, object]] = [
    ("revu_content_type", "블로그"),
    ("revu_purchase_combine", "아니오"),
    ("revu_urgent", "아니오"),
    ("revu_title", ""),
    ("revu_subtitle", ""),
    ("revu_car", ""),
    ("revu_product", ""),
    ("revu_qty", ""),
    ("revu_recruit", 10),
    ("revu_titlekw", ""),
    ("revu_bodykw", ""),
    ("revu_url", ""),
    ("revu_mission_0", ""),
    ("revu_mission_1", ""),
    ("revu_mission_2", ""),
    ("revu_mgr_name", DEFAULT_MANAGER_NAME),
    ("revu_mgr_phone", DEFAULT_MANAGER_PHONE),
    ("revu_mgr_email", DEFAULT_MANAGER_EMAIL),
    ("revu_track_base", ""),
    ("revu_nt_source", DEFAULT_NT_SOURCE),
    ("revu_nt_medium", DEFAULT_NT_MEDIUM),
    ("revu_nt_detail", DEFAULT_NT_DETAIL),
    ("revu_nt_keyword", ""),
]

# 라디오 위젯 키 → 허용값. 불러온 값이 허용 밖이면 기본값으로(StreamlitAPIException 방지).
REVU_RADIO_OPTIONS: dict[str, tuple[str, ...]] = {
    "revu_content_type": ("블로그", "클립"),
    "revu_purchase_combine": ("아니오", "예"),
    "revu_urgent": ("아니오", "예"),
}

_RECRUIT_MIN, _RECRUIT_MAX = 1, 999


def revu_form_defaults() -> dict:
    """저장 대상 위젯 key → 기본값 dict(세션 초기화·복원 기준)."""
    return {k: v for k, v in REVU_SAVE_FIELDS}


def serialize_form(values: dict) -> str:
    """위젯 값 dict → JSON 문자열(form_version 포함). 알려진 필드만 저장(잉여 무시)."""
    payload: dict = {"form_version": FORM_VERSION}
    for key, default in REVU_SAVE_FIELDS:
        payload[key] = values.get(key, default)
    return json.dumps(payload, ensure_ascii=False, indent=2)


def deserialize_form(raw) -> tuple[dict, list[str]]:
    """JSON(str/bytes) → ({위젯key: 값}, 경고목록). ★손상·구버전이어도 죽지 않는다.

    파싱 실패 → ({}, [에러]). 누락 필드는 채우지 않고(기본값 유지) 경고. 라디오 허용값·
    모집인원 범위/정수는 보정해 StreamlitAPIException(위젯 값 불일치)을 예방한다."""
    warnings: list[str] = []
    if isinstance(raw, (bytes, bytearray)):
        try:
            raw = bytes(raw).decode("utf-8")
        except UnicodeDecodeError:
            return {}, ["파일을 읽을 수 없습니다(UTF-8 인코딩 아님)."]
    try:
        data = json.loads(raw)
    except (json.JSONDecodeError, TypeError, ValueError):
        return {}, ["JSON 형식이 아닙니다(손상되었거나 양식 파일이 아님)."]
    if not isinstance(data, dict):
        return {}, ["양식 파일 구조가 아닙니다(JSON 객체가 아님)."]

    ver = data.get("form_version")
    if ver != FORM_VERSION:
        warnings.append(
            f"양식 버전이 다릅니다(파일 {ver!r} ≠ 현재 {FORM_VERSION}). 가능한 필드만 채웁니다."
        )

    out: dict = {}
    missing = 0
    for key, default in REVU_SAVE_FIELDS:
        if key not in data:
            missing += 1
            continue
        val = data[key]
        if key == "revu_recruit":
            try:
                val = max(_RECRUIT_MIN, min(_RECRUIT_MAX, int(val)))
            except (ValueError, TypeError):
                warnings.append("모집인원 값이 올바르지 않아 기본값으로 복원합니다.")
                val = default
        elif key in REVU_RADIO_OPTIONS:
            if val not in REVU_RADIO_OPTIONS[key]:
                warnings.append(f"{key} 값({val!r})이 유효하지 않아 기본값으로 복원합니다.")
                val = default
        else:
            val = default if val is None else str(val)
        out[key] = val

    if missing:
        warnings.append(f"일부 필드가 없어 기본값을 유지합니다({missing}개).")
    return out, warnings


def save_filename_json(values: dict) -> str:
    """저장 파일명: 양식_{제품명}_{차종}.json (빈 항목 생략, 알아보기 쉽게)."""
    product = str(values.get("revu_product", "") or "").strip().replace(" ", "")
    car = str(values.get("revu_car", "") or "").strip().replace(" ", "")
    parts = ["양식"]
    if product:
        parts.append(product)
    if car:
        parts.append(car)
    name = "_".join(parts).replace("/", "_").replace("\\", "_")
    return f"{name}.json"


# 허용 문자 패턴(공백은 별도 메시지로 먼저 잡는다).
_NT_ASCII_RE = re.compile(r"^[A-Za-z0-9._-]+$")          # source/medium
_NT_KO_RE = re.compile(r"^[가-힣ㄱ-ㅎㅏ-ㅣA-Za-z0-9._-]+$")  # detail/keyword(한글 허용)
_HANGUL_RE = re.compile(r"[가-힣ㄱ-ㅎㅏ-ㅣ]")


def validate_nt_param(name: str, value: str, *, required: bool, allow_korean: bool) -> str | None:
    """nt_ 파라미터 1개 검증. 통과면 None, 위반이면 한국어 경고 메시지 반환.

    검사 순서: 필수 누락 → 공백 → 한글(ascii 전용일 때) → 기타 특수문자."""
    v = value or ""
    if not v:
        return f"{name}: 필수값입니다 — 비우면 네이버가 추적하지 못합니다." if required else None
    if any(ws in v for ws in (" ", "\t", "\n", "　")):
        return f"{name}: 공백은 사용할 수 없습니다."
    if not allow_korean and _HANGUL_RE.search(v):
        return f"{name}: 한글은 사용할 수 없습니다 (영문·숫자·-, _, . 만 가능)."
    pattern = _NT_KO_RE if allow_korean else _NT_ASCII_RE
    if not pattern.match(v):
        return f"{name}: 허용되지 않는 문자가 있습니다 (-, _, . 외 특수문자·공백 금지)."
    return None


def build_tracking_url(
    product_url: str,
    nt_source: str,
    nt_medium: str,
    nt_detail: str = "",
    nt_keyword: str = "",
) -> str:
    """추적 파라미터를 제품 URL 에 조립(검증 없음 — 순수 조립).

    ★"?" 중복 방지: 제품 URL 에 이미 "?"가 있으면 "&"로, 없으면 "?"로 이어붙인다.
    선택 파라미터(detail/keyword)는 값이 있을 때만 포함, 필수 2개는 항상 포함."""
    params = [("nt_source", nt_source), ("nt_medium", nt_medium)]
    if nt_detail:
        params.append(("nt_detail", nt_detail))
    if nt_keyword:
        params.append(("nt_keyword", nt_keyword))
    query = "&".join(f"{k}={v}" for k, v in params)
    sep = "&" if "?" in product_url else "?"
    return f"{product_url}{sep}{query}"


def assemble_tracking_url(
    product_url: str,
    nt_source: str,
    nt_medium: str,
    nt_detail: str = "",
    nt_keyword: str = "",
) -> tuple[str | None, list[str]]:
    """검증 + 조립. 반환 (url 또는 None, 경고 메시지 목록).

    위반이 하나라도 있으면 url=None(생성 차단) 으로 errors 를 채워 돌려준다."""
    errors: list[str] = []
    base = (product_url or "").strip()
    if not base:
        errors.append("제품 URL: 필수값입니다.")
    elif any(ws in base for ws in (" ", "\t", "\n", "　")):
        errors.append("제품 URL: 공백이 포함되어 있습니다.")

    for name, val, required, allow_ko in (
        ("nt_source", nt_source, True, False),
        ("nt_medium", nt_medium, True, False),
        ("nt_detail", nt_detail, False, True),
        ("nt_keyword", nt_keyword, False, True),
    ):
        err = validate_nt_param(name, val, required=required, allow_korean=allow_ko)
        if err:
            errors.append(err)

    if errors:
        return None, errors
    return build_tracking_url(base, nt_source, nt_medium, nt_detail, nt_keyword), []

# 가벼운 금지어 목록(완벽한 필터 아님 — 경고용). 질병명 위주.
# 타브랜드·연예인명은 자동 판별이 어려워 직접 확인하도록 안내만 한다.
BANNED_WORDS = [
    "암", "당뇨", "고혈압", "고지혈", "아토피", "비염", "천식", "탈모",
    "관절염", "디스크", "치매", "우울증", "불면증", "변비", "치질",
    "코로나", "독감", "염증", "통증완화", "면역력", "항암", "항균",
    "다이어트", "체지방감소", "혈압", "혈당",
]
# 절대적·과장 표현(양식 4-7 규칙: "최고"·"최상" 등 금지).
ABSOLUTE_WORDS = ["최고", "최상", "최저가", "유일", "100%", "완벽", "1위"]


@dataclass
class RevuFormData:
    """레뷰 양식 입력값. 비어 있는 값은 양식의 빈칸 그대로 둔다."""
    content_type: str = "블로그"          # "블로그" | "클립"
    purchase_combine: str = "아니오"       # 구매평 결합 "예" | "아니오"
    urgent: str = "아니오"                 # 긴급 진행 "예" | "아니오"
    campaign_title: str = ""              # 캠페인 제목(≤20자)
    campaign_subtitle: str = ""           # 캠페인 부제목(≤40자)
    car_model: str = ""                   # 차종(선택)
    product_name: str = ""                # 제품명
    provide_qty: str = ""                 # 제공수량
    recruit_count: int = 10               # 모집인원
    title_keywords: str = ""              # 제목키워드(콤마 구분, 1~3개)
    body_keywords: str = ""               # 본문키워드(콤마 구분, 3~5개)
    product_url: str = ""                 # 제품링크 URL
    missions: list = field(default_factory=lambda: ["", "", ""])  # 미션 1·2·3
    manager_name: str = DEFAULT_MANAGER_NAME
    manager_phone: str = DEFAULT_MANAGER_PHONE
    manager_email: str = DEFAULT_MANAGER_EMAIL


def merge_keywords(existing: str, additions: list[str]) -> str:
    """기존 콤마구분 키워드 문자열에 새 키워드를 ★덧붙인다(덮어쓰지 않음).

    중복은 대소문자 무시로 제거하고, 등장 순서(기존 먼저)는 보존. ", " 로 join.
    키워드 추천 '칸에 추가' 버튼이 쓴다 — 기존 입력 손실 없이 누적."""
    out: list[str] = []
    seen: set[str] = set()
    for kw in [*existing.split(","), *additions]:
        k = kw.strip()
        if k and k.lower() not in seen:
            seen.add(k.lower())
            out.append(k)
    return ", ".join(out)


# ── 미션 자동 채움(실데이터 기반 5단 구조) ──────────────────────────────────
# 레뷰 상위 미션 분석(평균 600~800자)의 5단 구조를 3줄에 압축:
#   ①제목 형식 ②본문 글자수·키워드 반복 ③필수 언급 셀링포인트 ④사진/영상(+비교샷·AI금지)
#   ⑤차종 한정·구매처 링크·페널티 푸터.
# ★규칙: 모든 줄에 🔴정보형 단어(교체방법·셀프·장착·청소 등) ★절대 금지(구매형 지향).
# ★셀링포인트는 진짜 사실이라 사용자가 확인·수정한다(AI 생성 아님 — 허위 위험 회피).
# {cm}=차종, {prod}=제품명(없으면 "제품"), {suffix}=제목 접미, {benefit}=비교 포인트.

# 미션1은 제품 무관 공통(제목 형식 + 본문 글자수·키워드 반복).
_MISSION1 = (
    "제목은 「{cm} {prod} {suffix}」 형식으로 작성해주세요. "
    "제목 키워드를 본문에 12회 이상 자연스럽게 넣고, 본문은 2000~2500자로 작성해주세요."
)

# 제품군별 미션2·3 + 기본 셀링포인트 + 각도표[(key,label,suffix,benefit)].
_MISSION_BLOCKS = {
    "airfilter": {
        "selling": "활성탄 흡착 탈취 · 미세먼지/초미세먼지 차단 · 전 차종 호환 · 정품/가성비",
        "m2": (
            "다음 핵심 포인트를 반드시 언급해주세요 → [필수 언급(확인·수정): {SELLING}]. "
            "기존(순정/방치) 필터와 비교해 {benefit}{josa} 어떻게 달라졌는지 솔직한 실사용 "
            "후기로 적고, 다양한 차종 호환 가능 점도 함께 언급해주세요."
        ),
        "m3": (
            "사진 15장 이상 + 동작·사용 모습 동영상 2개(30초 이상)를 첨부하고, 기존에 쓰던 "
            "더러운 필터와 새 필터 비교 사진을 꼭 넣어주세요. 대표 사진은 연출 컷으로 예쁘게 "
            "지정하고 언박싱 사진은 2~3장만, AI로 생성한 이미지·영상은 사용 금지입니다. "
            "{cm} 차종만 신청해주시고, 포스팅 하단에 구매처 링크를 삽입해주세요. "
            "(미션 미준수 시 수정 요청이 있을 수 있습니다.)"
        ),
        "angles": [
            ("smell", "냄새", "냄새 잡는 활성탄 후기", "에어컨 냄새·퀴퀴함"),
            ("dust", "미세먼지", "미세먼지 차단 후기", "미세먼지 체감·실내 공기질"),
            ("fit", "호환·가성비", "호환 가성비 후기", "차량 적합성·가격 대비 만족도"),
            ("summer", "여름 냉방", "여름 에어컨 쾌적 후기", "여름철 냉방 시 쾌적함"),
        ],
    },
    "wiper": {
        "selling": "저소음·떨림 저감 · 발수 코팅 · 물자국/번짐 개선 · 규격(사이즈) 호환 · 가성비",
        "m2": (
            "다음 핵심 포인트를 반드시 언급해주세요 → [필수 언급(확인·수정): {SELLING}]. "
            "기존 와이퍼와 비교해 {benefit}{josa} 어떻게 달라졌는지 솔직한 실사용 후기로 적고, "
            "{cm}에 맞는 규격(사이즈) 호환 점도 함께 언급해주세요."
        ),
        "m3": (
            "와이퍼를 교체한 모습과 유리를 닦는 사진 7장 이상 + 동영상 1개 이상을 첨부하고, "
            "기존 와이퍼의 줄·번짐과 새 와이퍼가 깨끗하게 닦이는 비교 사진(또는 영상)을 꼭 "
            "넣어주세요. AI로 생성한 이미지·영상은 사용 금지입니다. {cm}에 맞는 규격으로 "
            "신청해주시고, 포스팅 하단에 구매처 링크를 삽입해주세요. "
            "(미션 미준수 시 수정 요청이 있을 수 있습니다.)"
        ),
        "angles": [
            ("noise", "소음·떨림", "소음 잡는 와이퍼 후기", "와이퍼 소음·떨림(채터링)"),
            ("water", "발수·시야", "발수 와이퍼 후기", "비 올 때 발수·시야 확보"),
            ("spec", "규격·호환", "규격 맞는 와이퍼 후기", "사이즈 적합성·가성비"),
            ("season", "장마·겨울", "장마철 와이퍼 후기", "빗물·성에 제거 성능"),
        ],
    },
    "generic": {
        "selling": "[제품 핵심 셀링포인트 2~3개 입력]",
        "m2": (
            "다음 핵심 포인트를 반드시 언급해주세요 → [필수 언급(확인·수정): {SELLING}]. "
            "기존 제품과 비교해 {benefit}{josa} 어떻게 달라졌는지 솔직한 실사용 후기로 적고, "
            "다양한 차종 호환 가능 점도 함께 언급해주세요."
        ),
        "m3": (
            "사진 15장 이상 + 동작·사용 모습 동영상 2개(30초 이상)를 첨부하고, 기존 제품과 "
            "새 제품 비교 사진을 꼭 넣어주세요. 대표 사진은 연출 컷으로 예쁘게 지정하고 "
            "언박싱 사진은 2~3장만, AI로 생성한 이미지·영상은 사용 금지입니다. "
            "{cm} 차종만 신청해주시고, 포스팅 하단에 구매처 링크를 삽입해주세요. "
            "(미션 미준수 시 수정 요청이 있을 수 있습니다.)"
        ),
        "angles": [
            ("value", "핵심가치", "실사용 후기", "사용 후 체감 변화"),
            ("price", "가성비", "가성비 후기", "가격 대비 만족도"),
        ],
    },
}


def _subject_josa(word: str) -> str:
    """주격 조사 보정 — 마지막 '한글' 음절에 받침 있으면 "이", 없으면 "가".

    한글이 없으면 "이". 괄호 등으로 끝나도 마지막 한글 음절을 찾아 판정
    (예: "채터링)" → '링'은 받침 있음 → "이")."""
    for ch in reversed(word):
        if "가" <= ch <= "힣":
            return "이" if (ord(ch) - 0xAC00) % 28 else "가"
    return "이"


def _product_kind(product_name: str) -> str:
    """제품명 → 미션 분기 키. 와이퍼 우선, 그다음 필터/에어컨/캐빈, 그 외 generic."""
    p = product_name or ""
    if "와이퍼" in p:
        return "wiper"
    if ("필터" in p) or ("에어컨" in p) or ("캐빈" in p):
        return "airfilter"
    return "generic"


def mission_angles(product_name: str) -> list[tuple[str, str]]:
    """제품군별 미션 '각도' 목록 [(key, label), ...] — 셀렉트박스용."""
    block = _MISSION_BLOCKS[_product_kind(product_name)]
    return [(key, label) for key, label, _suffix, _benefit in block["angles"]]


def mission_block(car_model: str, product_name: str, angle_key: str) -> list[str]:
    """차종·제품·각도로 리치 5단 구조 미션 3줄을 만든다(수정 가능). 차종 없으면 빈 3줄.

    angle_key 가 해당 제품군에 없으면 첫 각도로 폴백. 🔴정보형 단어는 ★쓰지 않는다."""
    cm = (car_model or "").strip()
    if not cm:
        return ["", "", ""]
    prod = (product_name or "").strip() or "제품"
    block = _MISSION_BLOCKS[_product_kind(product_name)]
    angle = next(
        (a for a in block["angles"] if a[0] == angle_key), block["angles"][0])
    _key, _label, suffix, benefit = angle
    fmt = dict(
        cm=cm, prod=prod, suffix=suffix, benefit=benefit,
        josa=_subject_josa(benefit), SELLING=block["selling"])
    return [
        _MISSION1.format(**fmt),
        block["m2"].format(**fmt),
        block["m3"].format(**fmt),
    ]


def default_mission_lines(car_model: str, product_name: str) -> list[str]:
    """하위호환 — 제품군의 ★첫 각도로 미션 3줄 생성. 차종 없으면 빈 3줄."""
    block = _MISSION_BLOCKS[_product_kind(product_name)]
    return mission_block(car_model, product_name, block["angles"][0][0])


def find_banned_words(*texts: str) -> list[str]:
    """입력 텍스트들에서 금지어/절대표현을 찾아 중복 없이 반환(경고용)."""
    blob = " ".join(t for t in texts if t)
    hits: list[str] = []
    for w in BANNED_WORDS + ABSOLUTE_WORDS:
        if w and w in blob and w not in hits:
            hits.append(w)
    return hits


def _set_para_text(paragraph, text: str) -> None:
    """문단 텍스트를 통째로 치환하되 첫 run 의 서식을 보존한다.

    첫 run 의 text 만 새 값으로 바꾸고 나머지 run 은 비워, 폰트·크기 등 서식을 유지.
    run 이 없으면(빈 문단) 새 run 을 추가한다."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _cell_para(doc, table_idx: int, row_idx: int, para_idx: int):
    """(표·행·문단) 좌표로 대상 문단을 안전하게 반환. 범위 밖이면 None."""
    try:
        cell = doc.tables[table_idx].rows[row_idx].cells[0]
        return cell.paragraphs[para_idx]
    except (IndexError, KeyError):
        return None


# Step1 드롭다운(콘텐츠 컨트롤 w:sdt)에 들어갈 선택 표시 텍스트(원본 dropDownList 항목과 동일).
_CONTENT_LABEL = {
    "블로그": "① [단독] 블로그 (Blog Only)",
    "클립": "② [단독] 클립 (Clip Only)",
}
_COMBINE_LABEL = {
    "예": "① 예 (구매평 결합 진행)",
    "아니오": "② 아니오 (미적용)",
}
_URGENT_LABEL = {
    "예": "① 예 (긴급 진행)",
    "아니오": "② 아니오 (미적용)",
}


def _unwrap_sdt_with_text(sdt, text: str) -> None:
    """드롭다운 위젯(w:sdt)을 ★제거하고 선택 텍스트 run 으로 대체한다.

    방식: sdtContent 안의 플레이스홀더("클릭하여 선택") run 텍스트를 선택값으로 바꾼 뒤,
    sdtContent 의 자식(run·permStart/End 등)을 부모 문단의 sdt 자리로 그대로 옮기고
    sdt 래퍼를 삭제한다. perm 마커는 짝을 유지(문서 무결성)."""
    from docx.oxml.ns import qn  # 지연 import — docx 생성 시점에만 필요.

    content = sdt.find(qn("w:sdtContent"))
    if content is None:
        return
    t_elems = content.findall(".//" + qn("w:t"))
    if t_elems:
        t_elems[0].text = text          # 플레이스홀더 → 선택값
        for t in t_elems[1:]:
            t.text = ""                  # 잔재 제거
    parent = sdt.getparent()
    insert_at = parent.index(sdt)
    for child in list(content):          # sdtContent 자식을 sdt 자리로 이동(순서 보존)
        parent.insert(insert_at, child)
        insert_at += 1
    parent.remove(sdt)                   # 위젯 래퍼 삭제


def _fill_step1_dropdowns(doc, data: RevuFormData) -> None:
    """표0 r1 의 드롭다운 3개(콘텐츠 타입·구매평 결합·긴급 진행)만 텍스트로 치환.

    ★정확히 그 3개만 타겟 — 표0 r1 셀 안의 w:sdt 만 순회하고, 각 sdt 의 listItem 값으로
    어떤 드롭다운인지 식별한다(다른 셀의 sdt·표12 등은 건드리지 않음)."""
    from docx.oxml.ns import qn  # 지연 import.

    try:
        tc = doc.tables[0].rows[1].cells[0]._tc
    except (IndexError, KeyError):
        return

    for sdt in tc.findall(".//" + qn("w:sdt")):
        values = " ".join(
            (li.get(qn("w:value")) or "") for li in sdt.findall(".//" + qn("w:listItem"))
        )
        if "Blog Only" in values or "Clip Only" in values:
            text = _CONTENT_LABEL.get(data.content_type)
        elif "구매평 결합 진행" in values:
            text = _COMBINE_LABEL.get(data.purchase_combine, _COMBINE_LABEL["아니오"])
        elif "긴급 진행" in values:
            text = _URGENT_LABEL.get(data.urgent, _URGENT_LABEL["아니오"])
        else:
            continue  # 식별 불가 sdt 는 안전하게 건너뜀(치환하지 않음).
        if text:
            _unwrap_sdt_with_text(sdt, text)


def fill_document(doc, data: RevuFormData) -> None:
    """로드된 Document 에 입력값을 채운다(in-place). 표12(사전안내문)는 건드리지 않는다.

    값이 비어 있는 항목은 양식 빈칸 그대로 둔다(없는 값을 만들어 넣지 않음)."""

    def put(t, r, p, text):
        para = _cell_para(doc, t, r, p)
        if para is not None:
            _set_para_text(para, text)

    # Step1 드롭다운 3개(콘텐츠 타입·구매평 결합·긴급 진행) — 위젯 제거 + 선택 텍스트 치환.
    # 라벨 run("2. 콘텐츠 타입 선택: " 등)은 그대로 두고, 드롭다운 자리에만 선택값을 박는다.
    _fill_step1_dropdowns(doc, data)

    # 담당자 — 표3 r0 p0~2
    put(3, 0, 0, f"성함: {data.manager_name}")
    put(3, 0, 1, f"연락처: {data.manager_phone}")
    put(3, 0, 2, f"이메일: {data.manager_email}")

    # 캠페인 제목/부제목 — 표4·5 r0 p0 (대괄호 안에)
    if data.campaign_title:
        put(4, 0, 0, f"[ {data.campaign_title} ]")
    if data.campaign_subtitle:
        put(5, 0, 0, f"[ {data.campaign_subtitle} ]")

    # 모집인원 — 표6 r0 p0 ("　　명" → "N명")
    if data.recruit_count is not None:
        put(6, 0, 0, f"{data.recruit_count}명")

    # 제품명/제공수량 — 표7 r0 p0·p1
    if data.product_name:
        put(7, 0, 0, f"제품명: {data.product_name}")
    # 제공수량 — 입력값에 수량 단위("2개" 등)까지 직접 포함하므로 원본의 "개"는 쓰지 않는다
    # (안 그러면 "...2개 개"로 중복됨). 입력값을 그대로 사용.
    if data.provide_qty:
        put(7, 0, 1, f"제공수량: {data.provide_qty}")

    # 미션 1·2·3 — 콘텐츠 타입에 따라 표8 r1(블로그) 또는 r3(클립)
    mission_row = 3 if data.content_type == "클립" else 1
    for i, m in enumerate(data.missions[:3]):
        if m:
            put(8, mission_row, i, f"{i + 1}. {m}")

    # 키워드 — 표10 r1 p0·p1
    if data.title_keywords:
        put(10, 1, 0, f"제목키워드 (1~3개) : {data.title_keywords}")
    if data.body_keywords:
        put(10, 1, 1, f"본문키워드 (3~5개) : {data.body_keywords}")

    # 제품링크 URL — 표11 r0 p0
    if data.product_url:
        put(11, 0, 0, f"링크입력: {data.product_url}")


def build_revu_docx(data: RevuFormData, template_path=None) -> bytes:
    """빈 양식을 로드 → 값 채움 → docx 바이트 반환(st.download_button 용)."""
    try:
        from docx import Document  # ★지연 import — docx 생성 시점에만 python-docx 필요.
    except ImportError as e:
        raise ImportError(
            "python-docx 가 설치돼 있지 않아 docx 를 생성할 수 없습니다. "
            "`pip install python-docx` (requirements.txt 에 python-docx>=1.1 포함)."
        ) from e
    path = pathlib.Path(template_path) if template_path else TEMPLATE_PATH
    if not path.exists():
        raise FileNotFoundError(
            f"레뷰 빈 양식 템플릿을 찾을 수 없습니다: {path}\n"
            "templates/revu_basic_template.docx 파일이 있는지 확인하세요."
        )
    doc = Document(str(path))
    fill_document(doc, data)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def suggest_filename(data: RevuFormData) -> str:
    """다운로드 파일명: 레뷰_{제품명}_{차종}.docx (빈 항목은 생략)."""
    parts = ["레뷰"]
    if data.product_name.strip():
        parts.append(data.product_name.strip())
    if data.car_model.strip():
        parts.append(data.car_model.strip())
    name = "_".join(parts).replace("/", "_").replace("\\", "_")
    return f"{name}.docx"
