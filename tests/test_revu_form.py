"""
test_revu_form.py — 레뷰 체험단 양식 docx 생성 회귀 테스트.

빈 양식(templates/revu_basic_template.docx)을 로드해 값을 채운 뒤, 결과 docx 를
다시 열어 ① 올바른 표 셀에 값이 들어갔는지 ② 사전안내문(표12)과 표 구조가
100% 보존됐는지 ③ 빈 값은 빈칸으로 남는지 ④ 미션 라우팅·금지어·파일명을 검증한다.
네트워크/Streamlit 없이 순수 함수만 다룬다.
"""

from __future__ import annotations

import io
import json
import subprocess
import sys
import textwrap

import pytest
from docx import Document
from docx.oxml.ns import qn

from src.revu_form import (
    FORM_VERSION,
    REVU_SAVE_FIELDS,
    SUBTITLE_MAX,
    TEMPLATE_PATH,
    TITLE_MAX,
    RevuFormData,
    assemble_tracking_url,
    build_revu_docx,
    build_tracking_url,
    default_mission_lines,
    deserialize_form,
    find_banned_words,
    merge_keywords,
    revu_form_defaults,
    save_filename_json,
    serialize_form,
    suggest_filename,
)
from src.core.keyword_reco import (
    extract_title_keywords,
    partition_banned,
    recommend_blog_keywords,
    recommend_keywords,
)

pytestmark = pytest.mark.skipif(
    not TEMPLATE_PATH.exists(),
    reason="빈 양식 템플릿(templates/revu_basic_template.docx)이 없으면 건너뜀",
)


def _build(**kw):
    """입력값으로 docx 를 만들고 다시 연 Document 를 반환."""
    data = RevuFormData(**kw)
    raw = build_revu_docx(data)
    return Document(io.BytesIO(raw))


def _cell(doc, t, r):
    return doc.tables[t].rows[r].cells[0].text


def _all_text(doc):
    """문서 전체 w:t 텍스트(merge) — sdt 잔재 검사용."""
    return "".join(t.text or "" for t in doc.element.body.findall(".//" + qn("w:t")))


def _n_sdt(doc):
    return len(doc.element.body.findall(".//" + qn("w:sdt")))


# ── 표 구조·사전안내문 보존 ──────────────────────────────────────────────────

def test_template_structure_unchanged():
    """생성 docx 의 표 개수·표12 행수가 원본과 동일해야 한다."""
    doc = _build(product_name="테스트")
    assert len(doc.tables) == 13
    assert len(doc.tables[12].rows) == 29


def test_pre_notice_table_byte_identical():
    """표12(사전안내문)의 모든 셀 텍스트가 원본과 글자 단위로 동일해야 한다(불변)."""
    original = Document(str(TEMPLATE_PATH))
    filled = _build(
        product_name="에어컨필터",
        title_keywords="키워드",
        campaign_title="제목",
    )
    orig_t12 = original.tables[12]
    fill_t12 = filled.tables[12]
    assert len(orig_t12.rows) == len(fill_t12.rows)
    for ri in range(len(orig_t12.rows)):
        o = [c.text for c in orig_t12.rows[ri].cells]
        f = [c.text for c in fill_t12.rows[ri].cells]
        assert o == f, f"표12 r{ri} 사전안내문이 변경됨"


# ── 셀별 값 배치 ─────────────────────────────────────────────────────────────

def test_values_land_in_correct_cells():
    doc = _build(
        content_type="블로그",
        campaign_title="EV5 에어컨필터",
        campaign_subtitle="여름철 차량 공기질 관리",
        product_name="파이널 에어컨필터",
        provide_qty="EV5 전용 P17 2개",
        recruit_count=15,
        title_keywords="에어컨필터교체",
        body_keywords="캐빈필터, 활성탄필터",
        product_url="https://brand.naver.com/dgk/products/1",
        missions=["미션1", "미션2", "미션3"],
        manager_name="박민우",
        manager_phone="010-3924-1155",
        manager_email="dgkorea93@naver.com",
    )
    assert "블로그" in _cell(doc, 0, 1)
    assert _cell(doc, 3, 0) == "성함: 박민우\n연락처: 010-3924-1155\n이메일: dgkorea93@naver.com"
    assert _cell(doc, 4, 0) == "[ EV5 에어컨필터 ]"
    assert _cell(doc, 5, 0) == "[ 여름철 차량 공기질 관리 ]"
    assert _cell(doc, 6, 0) == "15명"
    assert "제품명: 파이널 에어컨필터" in _cell(doc, 7, 0)
    assert "제공수량: EV5 전용 P17 2개" in _cell(doc, 7, 0)
    assert _cell(doc, 8, 1) == "1. 미션1\n2. 미션2\n3. 미션3"
    assert "제목키워드 (1~3개) : 에어컨필터교체" in _cell(doc, 10, 1)
    assert "본문키워드 (3~5개) : 캐빈필터, 활성탄필터" in _cell(doc, 10, 1)
    assert _cell(doc, 11, 0) == "링크입력: https://brand.naver.com/dgk/products/1"


def test_provide_qty_no_duplicate_gae():
    """제공수량은 입력값에 단위("2개")까지 포함 → 원본 "개"가 또 붙어 "...2개 개"가
    되면 안 된다(입력값 그대로, 끝이 "2개")."""
    doc = _build(provide_qty="EV5 에어컨필터 P17 2개")
    line = [p for p in _cell(doc, 7, 0).split("\n") if p.startswith("제공수량")][0]
    assert line == "제공수량: EV5 에어컨필터 P17 2개"
    assert "2개 개" not in line
    assert "개 개" not in _cell(doc, 7, 0)


def test_clip_missions_route_to_clip_row():
    """콘텐츠=클립이면 미션은 표8 r3(클립)에, 블로그 행(r1)은 비어 있어야 한다."""
    doc = _build(content_type="클립", missions=["가", "나", "다"])
    assert _cell(doc, 8, 3).startswith("1. 가\n2. 나\n3. 다")
    assert _cell(doc, 8, 1) == "1. \n2. \n3. "


def test_blog_missions_leave_clip_row_untouched():
    """콘텐츠=블로그면 클립 미션 행(r3)의 상품ID/해시태그 블록이 보존돼야 한다."""
    doc = _build(content_type="블로그", missions=["가", "나", "다"])
    clip = _cell(doc, 8, 3)
    assert clip.startswith("1. \n2. \n3. ")
    assert "상품ID" in clip and "해시태그" in clip


def test_empty_values_leave_blanks():
    """비운 항목은 양식 빈칸 그대로(없는 값을 만들어 넣지 않음)."""
    doc = _build()  # 전부 기본/빈값
    assert _cell(doc, 4, 0).strip() == "[      ]"   # 제목 빈칸 유지
    assert _cell(doc, 5, 0).strip() == "[      ]"   # 부제목 빈칸 유지
    assert _cell(doc, 11, 0).strip() == "링크입력:"  # 링크 라벨만
    # 제품명/제공수량 라벨만 남음
    assert "제품명:" in _cell(doc, 7, 0)


# ── Step1 드롭다운(w:sdt) 위젯 제거 + 선택 텍스트 치환 ────────────────────────

def test_dropdowns_replaced_with_selected_text_blog():
    """콘텐츠=블로그·구매평=예·긴급=아니오 → 표0 r1 에 선택 텍스트가 박힌다."""
    doc = _build(content_type="블로그", purchase_combine="예", urgent="아니오")
    cell = _cell(doc, 0, 1)
    assert "① [단독] 블로그 (Blog Only)" in cell
    assert "① 예 (구매평 결합 진행)" in cell
    assert "② 아니오 (미적용)" in cell           # 긴급=아니오


def test_dropdowns_replaced_with_selected_text_clip():
    """콘텐츠=클립·구매평=아니오·긴급=예."""
    doc = _build(content_type="클립", purchase_combine="아니오", urgent="예")
    cell = _cell(doc, 0, 1)
    assert "② [단독] 클립 (Clip Only)" in cell
    assert "① 예 (긴급 진행)" in cell
    assert "② 아니오 (미적용)" in cell           # 구매평=아니오


def test_dropdown_placeholder_residue_removed():
    """★'클릭하여 선택' 플레이스홀더 잔재가 없어야 한다(3개 드롭다운 모두 치환)."""
    doc = _build(content_type="블로그", purchase_combine="예", urgent="예")
    assert "클릭하여 선택" not in _all_text(doc)


def test_dropdown_widgets_removed_but_mission_sdt_kept():
    """표0 드롭다운 3개(w:sdt)는 제거(위젯→텍스트), 미션 셀의 4번째 sdt 는 보존."""
    orig = Document(str(TEMPLATE_PATH))
    assert _n_sdt(orig) == 4                       # 원본: 드롭다운3 + 미션1
    doc = _build(content_type="블로그")
    assert _n_sdt(doc) == 1                         # 표0 3개 제거, 미션 1개 유지
    # 미션 sdt 의 안내 플레이스홀더는 그대로 보존
    assert "클릭하여 항목을 선택하세요" in _all_text(doc)


def test_dropdown_replacement_keeps_docx_valid():
    """치환 후에도 docx 무손상 — zip 정상·document.xml well-formed·perm 마커 짝 유지."""
    import io as _io
    import zipfile
    from lxml import etree

    raw = build_revu_docx(RevuFormData(content_type="클립", purchase_combine="예", urgent="예"))
    z = zipfile.ZipFile(_io.BytesIO(raw))
    assert z.testzip() is None                      # zip 손상 없음
    etree.fromstring(z.read("word/document.xml"))   # XML well-formed(파싱 실패 시 예외)

    doc = Document(_io.BytesIO(raw))
    body = doc.element.body
    ps = len(body.findall(".//" + qn("w:permStart")))
    pe = len(body.findall(".//" + qn("w:permEnd")))
    assert ps == pe                                 # perm 마커 균형(문서 무결성)
    assert len(doc.tables) == 13
    assert len(doc.tables[12].rows) == 29           # 사전안내문 보존


def test_dropdown_defaults_are_no():
    """구매평·긴급 기본값은 '아니오' → 양쪽 다 '② 아니오 (미적용)'."""
    doc = _build(content_type="블로그")             # purchase_combine·urgent 기본
    cell = _cell(doc, 0, 1)
    assert cell.count("② 아니오 (미적용)") == 2


# ── 보조 함수 ────────────────────────────────────────────────────────────────

def test_char_limits():
    assert TITLE_MAX == 20
    assert SUBTITLE_MAX == 40


def test_find_banned_words_detects_disease_and_absolute():
    hits = find_banned_words("이 제품은 비염에 최고입니다")
    assert "비염" in hits
    assert "최고" in hits


def test_find_banned_words_clean_returns_empty():
    assert find_banned_words("에어컨필터 교체 후기", "캐빈필터 활성탄") == []


def test_default_mission_lines_with_car_model():
    lines = default_mission_lines("EV5", "에어컨필터")
    assert len(lines) == 3
    assert "EV5" in lines[0]
    assert "에어컨필터" in lines[0]


def test_default_mission_lines_without_car_model_blank():
    assert default_mission_lines("", "에어컨필터") == ["", "", ""]


def test_suggest_filename():
    data = RevuFormData(product_name="파이널 에어컨필터", car_model="EV5")
    assert suggest_filename(data) == "레뷰_파이널 에어컨필터_EV5.docx"


def test_suggest_filename_no_car_model():
    data = RevuFormData(product_name="와이퍼")
    assert suggest_filename(data) == "레뷰_와이퍼.docx"


# ── 네이버 유입 추적 URL(nt_) ────────────────────────────────────────────────

def test_tracking_no_question_mark_starts_with_question():
    """① "?" 없는 URL → "?"로 시작."""
    url = build_tracking_url("https://m.site.naver.com/2aAvQ", "naver.blog", "social")
    assert url == "https://m.site.naver.com/2aAvQ?nt_source=naver.blog&nt_medium=social"
    assert url.count("?") == 1


def test_tracking_existing_question_mark_uses_ampersand():
    """② 이미 "?" 있는 URL → "&"로 이어붙임("?" 중복 안 됨)."""
    url = build_tracking_url("https://store.naver.com/p?foo=1", "naver.blog", "social")
    assert url == "https://store.naver.com/p?foo=1&nt_source=naver.blog&nt_medium=social"
    assert url.count("?") == 1  # ★"?" 중복 방지


def test_tracking_korean_in_source_blocked():
    """③ nt_source 에 한글 → 경고 + 생성 차단(url None)."""
    url, errors = assemble_tracking_url(
        "https://m.site.naver.com/2aAvQ", "네이버블로그", "social")
    assert url is None
    assert any("한글" in e for e in errors)


def test_tracking_space_blocked():
    """④ 값에 공백 → 경고."""
    url, errors = assemble_tracking_url(
        "https://m.site.naver.com/2aAvQ", "naver blog", "social")
    assert url is None
    assert any("공백" in e for e in errors)


def test_tracking_required_empty_blocked():
    """⑤ 필수값(nt_source/nt_medium) 비면 차단."""
    url, errors = assemble_tracking_url(
        "https://m.site.naver.com/2aAvQ", "", "social")
    assert url is None
    assert any("nt_source" in e and "필수" in e for e in errors)


def test_tracking_optional_empty_excluded_required_kept():
    """⑥ 선택값 비면 그 파라미터 제외, 필수 2개는 유지."""
    url, errors = assemble_tracking_url(
        "https://m.site.naver.com/2aAvQ", "naver.blog", "social",
        nt_detail="", nt_keyword="")
    assert errors == []
    assert "nt_source=naver.blog" in url
    assert "nt_medium=social" in url
    assert "nt_detail" not in url
    assert "nt_keyword" not in url


def test_tracking_optional_present_included_with_korean():
    """선택값(한글 허용)이 있으면 포함된다 — nt_keyword 한글 OK."""
    url, errors = assemble_tracking_url(
        "https://m.site.naver.com/2aAvQ", "naver.blog", "social",
        nt_detail="revu", nt_keyword="EV5에어컨필터")
    assert errors == []
    assert "nt_detail=revu" in url
    assert "nt_keyword=EV5에어컨필터" in url


def test_tracking_disallowed_special_char_blocked():
    """③ 보강: 허용 외 특수문자(/,=,# 등)가 값에 있으면 차단."""
    url, errors = assemble_tracking_url(
        "https://m.site.naver.com/2aAvQ", "naver.blog", "soc=ial")
    assert url is None
    assert any("특수문자" in e for e in errors)


def test_tracking_product_url_empty_blocked():
    """제품 URL 자체가 비면 차단."""
    url, errors = assemble_tracking_url("", "naver.blog", "social")
    assert url is None
    assert any("제품 URL" in e and "필수" in e for e in errors)


# ── 키워드 추천(연관키워드 + 검색량) ─────────────────────────────────────────

def _fake_keywordstool(seed):
    """검색광고 키워드도구 응답 대역 — relKeyword + 월검색량 필드."""
    return [
        {"relKeyword": "EV5 에어컨필터 교체", "monthlyPcQcCnt": 500, "monthlyMobileQcCnt": 700},
        {"relKeyword": "EV5 에어컨필터 냄새", "monthlyPcQcCnt": 300, "monthlyMobileQcCnt": 500},
        {"relKeyword": "EV5 에어컨필터", "monthlyPcQcCnt": 100, "monthlyMobileQcCnt": 50},
        # 검색량 0(원래 '<10') → 제외돼야 함
        {"relKeyword": "극소키워드", "monthlyPcQcCnt": "< 10", "monthlyMobileQcCnt": 0},
        # 중복 relKeyword → dedupe
        {"relKeyword": "EV5 에어컨필터 교체", "monthlyPcQcCnt": 500, "monthlyMobileQcCnt": 700},
    ]


def test_recommend_keywords_returns_keyword_and_volume_sorted():
    """연관키워드+검색량 수집, 검색량 내림차순, 0·중복 제외."""
    out = recommend_keywords("EV5 에어컨필터", request_fn=_fake_keywordstool)
    assert out == [
        ("EV5 에어컨필터 교체", 1200),
        ("EV5 에어컨필터 냄새", 800),
        ("EV5 에어컨필터", 150),
    ]
    assert "극소키워드" not in [k for k, _ in out]  # 검색량 0 제외


def test_recommend_keywords_empty_seed():
    """빈 검색어 → 빈 리스트(네트워크 호출 없음)."""
    assert recommend_keywords("", request_fn=_fake_keywordstool) == []
    assert recommend_keywords("   ", request_fn=_fake_keywordstool) == []


def test_recommend_keywords_limit():
    """limit 으로 상위 N개만."""
    out = recommend_keywords("EV5 에어컨필터", request_fn=_fake_keywordstool, limit=2)
    assert len(out) == 2
    assert out[0] == ("EV5 에어컨필터 교체", 1200)


def test_partition_banned_excludes_disease_keywords():
    """추천 키워드 중 금지어(질병명) 포함 항목 분리."""
    pairs = [("에어컨필터 교체", 1200), ("비염 완화 필터", 300), ("탈모 방지", 100)]
    clean, excluded = partition_banned(pairs)
    assert ("에어컨필터 교체", 1200) in clean
    assert "비염 완화 필터" in excluded
    assert "탈모 방지" in excluded
    assert all("비염" not in k and "탈모" not in k for k, _ in clean)


# ── 블로그 제목 기반 키워드 추천 ─────────────────────────────────────────────

# 실제 블로그 검색 title 모양(<b> 강조 태그·HTML 엔티티 포함). 신차라 연관어는 빈약해도
# 제목엔 교체·교체방법·냄새·셀프·캐빈필터 같은 실제 표현이 풍부하게 등장한다.
_FAKE_BLOG_TITLES = [
    "기아 <b>EV5</b> 에어컨필터 교체 방법 셀프로 간단하게",
    "EV5 캐빈필터 교체주기 알아봤어요",
    "전기차 EV5 에어컨필터 냄새 제거 후기",
    "EV5 에어컨필터 셀프 교체 비용 정리",
    "EV5 에어컨필터 교체 했어요 (캐빈필터 추천)",
    "EV5 에어컨필터 교체 후기 &amp; 냄새 잡기",
]


def test_extract_title_keywords_frequency_and_seed_excluded():
    """제목에서 빈도순 키워드 추출 — 검색어 토큰(EV5·에어컨필터)은 제외, 실제 표현 확보."""
    kws = extract_title_keywords(_FAKE_BLOG_TITLES, "EV5 에어컨필터")
    words = [k for k, _ in kws]
    # 검색어 토큰은 빠진다
    assert "EV5" not in words and "에어컨필터" not in words
    # 신차 연관어가 빈약해도 제목에서 실제 키워드가 풍부하게 나온다
    assert "교체" in words
    assert "캐빈필터" in words
    assert "냄새" in words
    # 빈도순(교체가 가장 많은 제목에 등장) — 내림차순 정렬 확인
    counts = [c for _, c in kws]
    assert counts == sorted(counts, reverse=True)
    kw_map = dict(kws)
    assert kw_map["교체"] >= kw_map["캐빈필터"]


def test_extract_title_keywords_drops_stopwords_short_and_numbers():
    """불용어(알아봤어요·했어요)·1글자·숫자만 토큰은 제외된다."""
    titles = ["에어컨필터 교체 알아봤어요 2024 의 A 후기"]
    kws = extract_title_keywords(titles, "에어컨필터")
    words = [k for k, _ in kws]
    assert "알아봤어요" not in words   # 불용어
    assert "2024" not in words          # 숫자만
    assert "A" not in words and "의" not in words  # 1글자
    assert "교체" in words and "후기" in words


def test_extract_title_keywords_excludes_seed_compact_form():
    """검색어를 띄어 입력해도 블로그가 붙여 쓴 결합형(EV5에어컨필터)은 제외된다."""
    titles = ["EV5에어컨필터 교체 후기", "EV5에어컨필터 냄새"]
    words = [k for k, _ in extract_title_keywords(titles, "EV5 에어컨필터")]
    assert "EV5에어컨필터" not in words   # 검색어 결합형 = 노이즈, 제외
    assert "교체" in words and "냄새" in words


def test_extract_title_keywords_counts_per_title_once():
    """같은 단어가 한 제목에 여러 번 나와도 제목 단위 1회만 센다(스팸 과대평가 방지)."""
    titles = ["교체 교체 교체 에어컨필터", "교체 에어컨필터"]
    kws = dict(extract_title_keywords(titles, "에어컨필터"))
    assert kws["교체"] == 2   # 제목 2개에 등장 → 2 (총 4회가 아님)


def test_recommend_blog_keywords_with_injected_titles():
    """titles 직접 주입(fetch_blog_titles 가 이미 HTML 정제한 형태) → 키워드 추출 +
    원문 제목 그대로 보존 + 금지어 분리(네트워크 없음)."""
    clean_titles = [
        "기아 EV5 에어컨필터 교체 방법 셀프로 간단하게",
        "EV5 캐빈필터 교체주기 알아봤어요",
        "전기차 EV5 에어컨필터 냄새 제거 후기",
    ]
    out = recommend_blog_keywords("EV5 에어컨필터", titles=clean_titles)
    assert out["titles"] == clean_titles   # 원문 제목 그대로 보존(참고용)
    words = [k for k, _ in out["keywords"]]
    assert "교체" in words and "캐빈필터" in words and "냄새" in words


def test_recommend_blog_keywords_excludes_banned():
    """블로그 제목에서 뽑은 키워드 중 금지어(질병명)는 excluded 로 분리."""
    titles = ["에어컨필터 교체 비염 후기", "에어컨필터 냄새 교체"]
    out = recommend_blog_keywords("에어컨필터", titles=titles)
    words = [k for k, _ in out["keywords"]]
    assert "비염" in out["excluded"]
    assert "비염" not in words
    assert "교체" in words


def test_recommend_blog_keywords_empty_seed():
    """빈 검색어 → 빈 결과(네트워크 호출 없음)."""
    out = recommend_blog_keywords("", titles=_FAKE_BLOG_TITLES)
    assert out == {"keywords": [], "titles": [], "excluded": []}


def test_recommend_blog_keywords_uses_titles_fn_when_no_titles():
    """titles 미지정 → titles_fn(seed) 으로 가져온다(라이브 경로 대역)."""
    captured = {}

    def _fake_titles_fn(seed):
        captured["seed"] = seed
        return ["에어컨필터 교체 후기", "에어컨필터 냄새 제거"]

    out = recommend_blog_keywords("에어컨필터", titles_fn=_fake_titles_fn)
    assert captured["seed"] == "에어컨필터"
    assert "교체" in [k for k, _ in out["keywords"]]


def test_recommend_blog_keywords_propagates_fetch_error():
    """블로그 호출 실패는 예외로 전파 → 호출부(app)가 경고만 띄우고 연관어는 살린다."""
    def _boom(seed):
        raise RuntimeError("429 Too Many Requests")

    with pytest.raises(RuntimeError):
        recommend_blog_keywords("에어컨필터", titles_fn=_boom)


# ── 블로그 제목: 제품 관련성 필터(차량 일반 글 제외) ─────────────────────────

# EV5 일반 글(보조금·연비 등)이 섞여 반환되는 실제 상황 재현.
_MIXED_BLOG_TITLES = [
    "기아 EV5 에어컨필터 교체주기 냄새 활성탄 후기",   # 제품 글 ✔
    "EV5 캐빈필터 셀프 교체 방법",                      # 제품 글 ✔
    "EV5 보조금 연비 풀체인지 총정리",                  # 무관 글 ✘
    "EV5 주행거리 제원표 모의견적 유지비",             # 무관 글 ✘
    "기아 EV5 GT 트림 출고 계약 후기",                  # 무관 글 ✘
]


def test_blog_extracts_only_from_product_titles():
    """제품(에어컨필터)어가 든 제목만 추출 — EV5 일반 글의 무관어는 안 나온다."""
    out = recommend_blog_keywords("EV5 에어컨필터", titles=_MIXED_BLOG_TITLES)
    words = [k for k, _ in out["keywords"]]
    for noise in ["보조금", "연비", "풀체인지", "주행거리", "제원표",
                  "모의견적", "유지비", "트림", "출고", "계약"]:
        assert noise not in words, f"무관어가 추출됨: {noise}"
    # 제품 관련어는 정상 추출
    assert "교체주기" in words
    assert "냄새" in words
    assert "활성탄" in words
    assert "캐빈필터" in words


def test_blog_excludes_vehicle_general_and_brand_words():
    """제품 제목 안에 섞인 차량 일반어·브랜드는 제외, 제품 관련어는 유지."""
    titles = ["에어컨필터 교체 기아 보조금 연비 활성탄 헤파"]
    out = recommend_blog_keywords("에어컨필터", titles=titles)
    words = [k for k, _ in out["keywords"]]
    assert "기아" not in words and "보조금" not in words and "연비" not in words
    # ★제품 관련어는 살아남는다
    assert "활성탄" in words and "헤파" in words and "교체" in words


def test_blog_wiper_only_product_titles():
    """와이퍼도 동일 — 와이퍼 무관어(연비·유지비) 제외, 와이퍼 관련어만."""
    titles = [
        "쏘렌토 와이퍼 교체주기 사이즈 발수 후기",   # 제품 글 ✔
        "쏘렌토 연비 유지비 보조금 신형 총정리",      # 무관 글 ✘
    ]
    out = recommend_blog_keywords("쏘렌토 와이퍼", titles=titles)
    words = [k for k, _ in out["keywords"]]
    assert "교체주기" in words and "사이즈" in words and "발수" in words
    assert "연비" not in words and "유지비" not in words and "보조금" not in words


def test_blog_counts_product_titles():
    """결과 dict 에 제품 관련 제목 수/전체 제목 수가 담긴다(앱 안내문용)."""
    out = recommend_blog_keywords("EV5 에어컨필터", titles=_MIXED_BLOG_TITLES)
    assert out["n_total_titles"] == 5
    assert out["n_product_titles"] == 2


def test_blog_fallback_when_no_product_titles():
    """제품어가 든 제목이 하나도 없어도 죽지 않고 빈 결과 + 원문은 보존."""
    titles = ["EV5 보조금 연비 총정리", "EV5 풀체인지 주행거리"]
    out = recommend_blog_keywords("EV5 에어컨필터", titles=titles)
    assert out["keywords"] == []          # 추출 0(노이즈 안 끌어옴)
    assert out["n_product_titles"] == 0
    assert out["titles"] == titles        # 원문 제목은 참고용으로 그대로 보존


def test_detect_product_synonyms():
    """검색어에서 제품군 식별 — 차종명만 있으면 None(필터 미적용)."""
    from src.core.keyword_reco import detect_product_synonyms

    af = detect_product_synonyms("EV5 에어컨필터")
    assert af and "캐빈필터" in af and "필터" in af
    wp = detect_product_synonyms("쏘렌토 와이퍼")
    assert wp and "블레이드" in wp
    assert detect_product_synonyms("EV5") is None       # 제품어 없음 → 필터 안 함
    assert detect_product_synonyms("") is None


# ── 블로그 제목 수집(HTML 정제) ──────────────────────────────────────────────

def test_fetch_blog_titles_strips_html_and_skips_empty():
    """fetch_blog_titles: <b> 태그·엔티티 제거, 빈 제목 스킵(http_get 주입)."""
    from src.core.teamp_mode import fetch_blog_titles

    class _Resp:
        status_code = 200
        headers: dict = {}

        def json(self):
            return {"items": [
                {"title": "EV5 <b>에어컨필터</b> 교체 &amp; 후기"},
                {"title": ""},   # 빈 제목 → 스킵
                {"title": "캐빈필터 교체주기"},
            ]}

        def raise_for_status(self):
            pass

    titles = fetch_blog_titles(
        "EV5 에어컨필터", "cid", "csec",
        http_get=lambda *a, **k: _Resp(), sleep_fn=lambda s: None)
    assert titles == ["EV5 에어컨필터 교체 & 후기", "캐빈필터 교체주기"]


# ── 키워드 칸 덧붙이기(merge) ────────────────────────────────────────────────

def test_merge_keywords_into_empty():
    assert merge_keywords("", ["에어컨필터교체", "에어컨필터냄새"]) == "에어컨필터교체, 에어컨필터냄새"


def test_merge_keywords_appends_not_overwrite():
    """기존 값에 덧붙임(덮어쓰지 않음)."""
    assert merge_keywords("기존키워드", ["새키워드"]) == "기존키워드, 새키워드"


def test_merge_keywords_dedupes():
    """중복(대소문자 무시) 자동 제거, 순서 보존."""
    assert merge_keywords("EV5, 에어컨필터", ["에어컨필터", "ev5", "냄새"]) == "EV5, 에어컨필터, 냄새"


def test_merge_keywords_no_additions_keeps_existing():
    assert merge_keywords("a, b", []) == "a, b"


# ── 배포본 회귀: python-docx 없어도 모듈 import·순수함수 동작 ────────────────

def test_revu_form_imports_without_python_docx():
    """★배포본 ImportError 회귀: python-docx 가 없어도 src.revu_form import 와
    순수 함수(merge_keywords/find_banned_words)는 동작해야 한다.

    docx 를 import 차단한 서브프로세스에서 검증 — 최상단 docx import 가 재도입되면
    이 테스트가 깨진다(app.py line 76 'from src.revu_form import' 동반 사망 방지)."""
    code = textwrap.dedent(
        """
        import builtins
        _real = builtins.__import__
        def _blocked(name, *a, **k):
            if name == "docx" or name.startswith("docx."):
                raise ModuleNotFoundError("No module named 'docx'")
            return _real(name, *a, **k)
        builtins.__import__ = _blocked

        from src.revu_form import merge_keywords, find_banned_words, assemble_tracking_url
        assert merge_keywords("a", ["b", "a"]) == "a, b"
        assert "비염" in find_banned_words("비염 완화")
        url, errors = assemble_tracking_url("https://x/p", "naver.blog", "social")
        assert url and not errors
        print("OK")
        """
    )
    proc = subprocess.run(
        [sys.executable, "-c", code],
        capture_output=True, text=True, cwd=str(TEMPLATE_PATH.parent.parent),
    )
    assert proc.returncode == 0, f"docx 없이 import 실패:\nSTDOUT={proc.stdout}\nSTDERR={proc.stderr}"
    assert "OK" in proc.stdout


def test_build_revu_docx_clear_error_without_python_docx():
    """python-docx 없을 때 build_revu_docx 는 명확한 ImportError 를 던진다(조용한 실패 금지)."""
    code = textwrap.dedent(
        """
        import builtins
        _real = builtins.__import__
        def _blocked(name, *a, **k):
            if name == "docx" or name.startswith("docx."):
                raise ModuleNotFoundError("No module named 'docx'")
            return _real(name, *a, **k)
        builtins.__import__ = _blocked

        from src.revu_form import build_revu_docx, RevuFormData
        try:
            build_revu_docx(RevuFormData(product_name="x"))
            print("NO_ERROR")
        except ImportError as e:
            assert "python-docx" in str(e)
            print("RAISED")
        """
    )
    proc = subprocess.run(
        [sys.executable, "-c", code],
        capture_output=True, text=True, cwd=str(TEMPLATE_PATH.parent.parent),
    )
    assert proc.returncode == 0, f"STDERR={proc.stderr}"
    assert "RAISED" in proc.stdout


# ── 양식 저장/불러오기(JSON) ─────────────────────────────────────────────────

_SAMPLE_VALUES = {
    "revu_content_type": "클립",
    "revu_purchase_combine": "예",
    "revu_urgent": "예",
    "revu_title": "EV5 에어컨필터 교체",
    "revu_subtitle": "여름철 차량 공기질 관리",
    "revu_car": "EV5",
    "revu_product": "파이널 에어컨필터",
    "revu_qty": "EV5 전용 P17 2개",
    "revu_recruit": 15,
    "revu_titlekw": "에어컨필터교체",
    "revu_bodykw": "캐빈필터, 활성탄필터",
    "revu_url": "https://brand.naver.com/dgk/products/1",
    "revu_mission_0": "미션1", "revu_mission_1": "미션2", "revu_mission_2": "미션3",
    "revu_mgr_name": "홍길동", "revu_mgr_phone": "010-0000-0000", "revu_mgr_email": "x@y.com",
    "revu_track_base": "https://m.site.naver.com/2aAvQ",
    "revu_nt_source": "naver.blog", "revu_nt_medium": "social",
    "revu_nt_detail": "revu", "revu_nt_keyword": "EV5에어컨필터",
}


def test_serialize_includes_all_fields_and_version():
    """저장 JSON 은 form_version + 모든 저장 필드를 포함한다."""
    js = serialize_form(_SAMPLE_VALUES)
    data = json.loads(js)
    assert data["form_version"] == FORM_VERSION
    for key, _ in REVU_SAVE_FIELDS:
        assert key in data, f"저장 누락: {key}"
    assert data["revu_recruit"] == 15
    assert data["revu_content_type"] == "클립"


def test_serialize_missing_value_uses_default():
    """값이 빠진 위젯은 기본값으로 저장(없는 값 생성 안 함)."""
    js = serialize_form({"revu_product": "와이퍼"})
    data = json.loads(js)
    assert data["revu_product"] == "와이퍼"
    assert data["revu_content_type"] == "블로그"     # 기본값
    assert data["revu_recruit"] == 10                # 기본값


def test_roundtrip_save_then_load_restores_values():
    """★왕복: 저장(JSON) → 불러오기 → 모든 값이 정확히 복원."""
    js = serialize_form(_SAMPLE_VALUES)
    restored, warnings = deserialize_form(js)
    assert warnings == []
    for key, _ in REVU_SAVE_FIELDS:
        assert restored[key] == _SAMPLE_VALUES[key], f"복원 불일치: {key}"


def test_roundtrip_via_bytes():
    """업로드 파일은 bytes — bytes 입력도 복원돼야 한다."""
    js = serialize_form(_SAMPLE_VALUES).encode("utf-8")
    restored, warnings = deserialize_form(js)
    assert restored["revu_content_type"] == "클립"
    assert restored["revu_recruit"] == 15


def test_deserialize_corrupt_json_does_not_crash():
    """손상 JSON → 예외 없이 ({}, [에러])."""
    restored, warnings = deserialize_form("{ not valid json ")
    assert restored == {}
    assert warnings and any("JSON" in w for w in warnings)


def test_deserialize_non_object_json():
    """JSON 이지만 객체가 아니면(리스트 등) 안전 처리."""
    restored, warnings = deserialize_form("[1, 2, 3]")
    assert restored == {}
    assert warnings


def test_deserialize_missing_fields_keeps_defaults_with_warning():
    """구버전·부분 파일: 있는 필드만 채우고 없는 필드는 건너뜀(경고)."""
    js = json.dumps({"form_version": FORM_VERSION, "revu_product": "와이퍼"})
    restored, warnings = deserialize_form(js)
    assert restored["revu_product"] == "와이퍼"
    assert "revu_content_type" not in restored      # 없는 필드는 안 채움(기본값 유지)
    assert any("일부 필드" in w for w in warnings)


def test_deserialize_invalid_radio_value_falls_back():
    """라디오 허용 밖 값 → 기본값으로 복원(+경고). StreamlitAPIException 예방."""
    js = json.dumps({"form_version": FORM_VERSION, "revu_content_type": "엉뚱한값"})
    restored, warnings = deserialize_form(js)
    assert restored["revu_content_type"] == "블로그"  # 기본값
    assert any("유효하지 않" in w for w in warnings)


def test_deserialize_recruit_clamped_and_coerced():
    """모집인원: 정수 보정 + 범위(1~999) 클램프(number_input 범위 위반 방지)."""
    js = json.dumps({"form_version": FORM_VERSION, "revu_recruit": 99999})
    restored, _ = deserialize_form(js)
    assert restored["revu_recruit"] == 999
    js2 = json.dumps({"form_version": FORM_VERSION, "revu_recruit": "abc"})
    restored2, warnings2 = deserialize_form(js2)
    assert restored2["revu_recruit"] == 10           # 기본값
    assert any("모집인원" in w for w in warnings2)


def test_deserialize_version_mismatch_warns_but_loads():
    """버전 불일치 → 경고하되 가능한 필드는 복원."""
    js = json.dumps({"form_version": 999, "revu_product": "와이퍼"})
    restored, warnings = deserialize_form(js)
    assert restored["revu_product"] == "와이퍼"
    assert any("버전" in w for w in warnings)


def test_save_filename_json():
    assert save_filename_json(_SAMPLE_VALUES) == "양식_파이널에어컨필터_EV5.json"
    assert save_filename_json({}) == "양식.json"


def test_revu_form_defaults_matches_save_fields():
    d = revu_form_defaults()
    assert set(d) == {k for k, _ in REVU_SAVE_FIELDS}
